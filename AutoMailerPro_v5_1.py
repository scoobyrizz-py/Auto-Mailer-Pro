
#!/usr/bin/env python3
"""
InsuranceMailer v1.0
Author: Kyle Padilla
Company: Jones Insurance Advisors, Inc.
Contact: scooby_rizz@proton.me

Description:
    Automated script to generate personalized letters, envelopes,
    and mailing labels for owner-occupied properties (personal) or businesses (commercial).

Usage:
    python InsuranceMailer_v1.0.py
    OR called from GUI with mode, file_path, content, and subject_line parameters.

Requirements:
    pandas, python-docx, fuzzywuzzy, python-Levenshtein
"""

__version__ = "1.0"
__author__ = "Kyle Padilla"
__company__ = "Jones Insurance Advisors, Inc."
__contact__ = "scooby_rizz@proton.me"

import os
import pandas as pd
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from fuzzywuzzy import fuzz
import csv

# === CONFIG ===
ZIP_LOOKUP_FILE = "zip_lookup.csv"
LOGO_PATH = "logo.png"

YOUR_NAME = "Brian Jones"
YOUR_TITLE = "Vice President"
YOUR_CO = "Jones Insurance Advisors, Inc"
YOUR_PHONE = "(772) 569-6802"
YOUR_EMAIL = "Brian@jonesia.com"
YOUR_ADDRESS = "3885 20th Street,\n Vero Beach, FL 32960"
YOUR_WEB = "www.jonesinsuranceadvisors.com"
YOUR_RETURN_ADDRESS = f"{YOUR_NAME}\n{YOUR_ADDRESS}"

# === ZIP TO CITY/STATE LOOKUP ===
zip_city_state = {}

def load_zip_lookup():
    global zip_city_state
    if not os.path.exists(ZIP_LOOKUP_FILE):
        print(f"❌ Missing ZIP lookup file: {ZIP_LOOKUP_FILE}")
        return
    df = pd.read_csv(ZIP_LOOKUP_FILE, dtype=str)
    for _, row in df.iterrows():
        zip_code = row['zip'].zfill(5)
        city_state = f"{row['city'].title()}, {row['state'].upper()}"
        zip_city_state[zip_code] = city_state

def zip_to_city_state(zip_code):
    zip_code = str(zip_code).zfill(5)
    city_state = zip_city_state.get(zip_code, "Indian River County, FL")
    return f"{city_state} {zip_code}"

# === CLEAN NAME ===
def clean_name(raw_name, mode):
    raw_name = str(raw_name)
    if mode == "personal":
        name_parts = [part.strip() for part in raw_name.split('||')]
        last_names = []
        first_names = []
        for part in name_parts:
            name_no_suffix = part.split('(')[0].strip()
            words = name_no_suffix.split()
            if len(words) >= 2:
                last_names.append(words[0].title())
                first_names.append(words[1].title())
            elif len(words) == 1:
                last_names.append("")
                first_names.append(words[0].title())
            else:
                last_names.append("")
                first_names.append("")
        unique_last_names = set([ln for ln in last_names if ln])
        if len(unique_last_names) == 1:
            last_name = unique_last_names.pop()
            combined_first_names = " & ".join([fn for fn in first_names if fn])
            full_name = f"{combined_first_names} {last_name}".strip()
        else:
            full_name = " & ".join(
                [f"{fn} {ln}".strip() for fn, ln in zip(first_names, last_names) if fn or ln]
            )
        return full_name or "Valued Customer"
    else:  # commercial
        return raw_name.title() or "Valued Business"

# === FILTERS ===
def is_owner_occupied(property_address, mailing_address):
    try:
        prop_addr = str(property_address).lower().strip()
        mailing_parts = str(mailing_address).split('|')
        for part in mailing_parts:
            if fuzz.partial_ratio(prop_addr, part.lower().strip()) > 85:
                return True
        return False
    except:
        return False

def is_valid_business(business_type):
    # Define valid business types (customize as needed)
    valid_types = ["Retail", "Office", "Restaurant", "Manufacturing", "Services"]
    return str(business_type).strip() in valid_types

# === ADD LETTER TO DOC ===
def add_letter_to_doc(doc, name, address, zip_code, sale_date, sale_price, content, mode, subject_line):
    today = datetime.now().strftime('%B %d, %Y')

    def add_compact_paragraph(text="", bold=False, space_before=0, space_after=2):
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(space_before)
        para.paragraph_format.space_after = Pt(space_after)
        run = para.add_run(text)
        run.bold = bold
        return para

    for _ in range(4):
        doc.add_paragraph()

    add_compact_paragraph(today, space_after=48)
    greeting = f"Dear {name},"
    add_compact_paragraph(greeting, space_after=24)

    # Use the provided subject line (bolded)
    add_compact_paragraph(subject_line, bold=True, space_after=10)

    city_state = zip_to_city_state(zip_code)
    county_name = "Indian River" if "County" in city_state else city_state.split("County")[0].strip()
    personalized_content = content.replace("[Name]", name).replace("[County]", county_name)

    doc.add_paragraph(personalized_content)

    if os.path.exists("signature_brian.png"):
        doc.add_picture("signature_brian.png", width=Inches(1.5), height=Inches(0.5))

    doc.add_paragraph(
        f"{YOUR_NAME}\n{YOUR_TITLE}\n{YOUR_EMAIL}\n{YOUR_PHONE}\n{YOUR_WEB}"
    )

    doc.add_page_break()

# === ADD ENVELOPE TO DOC ===
def add_envelope_to_doc(doc, name, address, zip_code):
    section = doc.add_section()
    section.page_width = Inches(9.5)
    section.page_height = Inches(4.125)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)

    sender = doc.add_paragraph(YOUR_RETURN_ADDRESS)
    sender.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    for run in sender.runs:
        run.font.size = Pt(10)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(40)

    recipient = doc.add_paragraph()
    recipient.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    name_run = recipient.add_run(f"{name}\n")
    name_run.bold = True
    name_run.font.size = Pt(14)
    addr_run = recipient.add_run(f"{address}\n{zip_to_city_state(zip_code)}")
    addr_run.font.size = Pt(14)

    doc.add_page_break()

# === CREATE LABELS DOC ===
def create_labels(label_data, labels_file):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.19)
    section.right_margin = Inches(0.19)

    labels_per_row = 3
    label_width = Inches(2.63)
    label_height = Inches(1.0)
    num_rows = (len(label_data) + labels_per_row - 1) // labels_per_row

    table = doc.add_table(rows=num_rows, cols=labels_per_row)
    table.autofit = False
    table.allow_autofit = False

    for col in table.columns:
        col.width = label_width
    for row in table.rows:
        row.height = label_height
        row.height_rule = 2

    idx = 0
    for row in table.rows:
        for cell in row.cells:
            if idx < len(label_data):
                lines = label_data[idx].split("\n")
                para = cell.paragraphs[0]
                para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                if lines:
                    name_run = para.add_run(lines[0] + "\n")
                    name_run.bold = True
                    name_run.font.size = Pt(10.5)
                for line in lines[1:]:
                    addr_run = para.add_run(line + "\n")
                    addr_run.font.size = Pt(10.5)
                idx += 1
            else:
                cell.text = ""

    doc.save(labels_file)
    print(f"✅ Mailing labels saved to: {labels_file}")

# === MAIN ===
def main(mode="personal", file_path="sales_data.xlsx", content=None, subject_line=""):
    if mode not in ["personal", "commercial"]:
        raise ValueError("Mode must be 'personal' or 'commercial'")
    if not subject_line:
        if mode == "personal":
            subject_line = "Homeowners Insurance Rates Are Finally on the Decline – Don’t Miss Out!"
        else:
            subject_line = "Protect Your Business with Tailored Insurance Solutions!"

    if content is None:
        if mode == "personal":
            content = (
                "For the first time in years, homeowners rates are coming down — and the savings could be significant.\n\n"
                "Recent legislative changes have boosted competition in Florida’s property insurance market, "
                "and many [County] County homeowners are already benefiting.\n\n"
                "Jones Insurance Advisors is a two-generation, family-owned independent agency located right here in Vero Beach. "
                "Our team of dedicated agents possess extensive knowledge of the intricacies of the local insurance market, "
                "and are excited to assist you in finding the most comprehensive and competitively priced insurance solutions.\n\n"
                "Call us today for a free, no-obligation quote, or visit our website below and complete a quote request, "
                "and one of our dedicated agents will reach out to you!\n\n"
                "We look forward to earning your business and providing you the personal, dedicated service you have come to "
                "expect by doing business locally.\n\n"
                "Warm Regards,"
            )
        else:
            content = (
                "Protecting your business is our priority at Jones Insurance Advisors.\n\n"
                "As a [County] County business, you need insurance solutions tailored to your unique needs. "
                "Our experienced team specializes in crafting comprehensive coverage plans for businesses like yours, "
                "ensuring protection against risks while keeping costs competitive.\n\n"
                "Jones Insurance Advisors, a family-owned agency in Vero Beach, is here to help. "
                "Contact us for a free consultation to discuss how we can safeguard your business.\n\n"
                "We look forward to partnering with you!\n\n"
                "Best Regards,"
            )

    timestamp = datetime.now().strftime("%m%d%y_%H%M%S")
    OUTPUT_DIR = f"{timestamp}_{mode.capitalize()}_Mailing_Campaign"
    LETTERS_FILE = os.path.join(OUTPUT_DIR, "all_letters.docx")
    ENVELOPES_FILE = os.path.join(OUTPUT_DIR, "all_envelopes.docx")
    LABELS_FILE = os.path.join(OUTPUT_DIR, "mailing_labels.docx")
    CRM_EXPORT_FILE = os.path.join(OUTPUT_DIR, f"crm_{mode}_occupied.csv")

    if not os.path.exists(OUTPUT_DIR):
        os.makedirs(OUTPUT_DIR)
        print(f"📁 Created output folder: {OUTPUT_DIR}")

    letters_doc = Document()
    envelopes_doc = Document()

    load_zip_lookup()

    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Excel file not found: {file_path}")

    try:
        df = pd.read_excel(file_path)
    except Exception as e:
        raise Exception(f"Failed to read Excel file: {e}")

    labels = []
    crm_rows = []

    for _, row in df.iterrows():
        try:
            if mode == "personal":
                name_key = 'Owner Name'
                filter_check = is_owner_occupied(row.get('Address', ''), row.get('Mailing Address', ''))
                filter_desc = "non-owner-occupied"
            else:  # commercial
                name_key = 'Business Name'
                filter_check = is_valid_business(row.get('Business Type', ''))
                filter_desc = "invalid business type"

            name = str(row.get(name_key, ''))
            if not name:
                print(f"⏭️ Skipping row with missing {name_key}")
                continue

            address = str(row.get('Address', '')).title().strip()
            zip_code = str(row.get('Site Zip Code', '')).strip()
            mailing_address = str(row.get('Mailing Address', '')).strip()
            sale_date_raw = str(row.get('Sale Date', '')).strip()
            sale_price_str = str(row.get('Sale Price', '')).replace('$', '').replace(',', '').strip()

            if not filter_check:
                print(f"⏭️ Skipping {filter_desc}: {name}")
                continue

            name = clean_name(name, mode)
            try:
                sale_price = float(sale_price_str) if sale_price_str else 0.0
            except ValueError:
                sale_price = 0.0
            try:
                sale_date = datetime.strptime(sale_date_raw, '%m/%d/%Y').strftime('%B %d, %Y') if sale_date_raw else "Unknown"
            except ValueError:
                sale_date = "Unknown"

            add_letter_to_doc(letters_doc, name, address, zip_code, sale_date, sale_price, content, mode, subject_line)
            add_envelope_to_doc(envelopes_doc, name, address, zip_code)

            label_text = f"{name}\n{address}\n{zip_to_city_state(zip_code)}"
            labels.append(label_text)

            crm_rows.append({
                'Name': name,
                'Address': address,
                'Zip': zip_code,
                'Sale Date': sale_date,
                'Sale Price': sale_price,
                'Email': '',
                'Phone': '',
                'Source': f"{mode.capitalize()} Anniversary Mailer-Sept-Oct"
            })

            print(f"✅ Processed: {name}")

        except Exception as e:
            print(f"⚠️ Skipped row due to error: {e}")

    if labels:
        create_labels(labels, LABELS_FILE)

    if crm_rows:
        keys = crm_rows[0].keys()
        with open(CRM_EXPORT_FILE, 'w', newline='', encoding='utf-8') as f:
            dict_writer = csv.DictWriter(f, keys)
            dict_writer.writeheader()
            dict_writer.writerows(crm_rows)
        print(f"📥 CRM-ready CSV saved to: {CRM_EXPORT_FILE}")

    letters_doc.save(LETTERS_FILE)
    envelopes_doc.save(ENVELOPES_FILE)
    print(f"📄 All letters saved to: {LETTERS_FILE}")
    print(f"✉️ All envelopes saved to: {ENVELOPES_FILE}")

def print_logo():
    logo = r"""
                      __/___             
             _____/______|           
     _______/_____\_______\_____     
     \              < < <       |    
      ~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~  
       Auto Mailer Pro v5.1
       Author: Kyle Padilla
       Last Updated: 08/09/2025
       
           Version: 05.01
    """
    print(logo)

if __name__ == "__main__":
    print_logo()
    main()
