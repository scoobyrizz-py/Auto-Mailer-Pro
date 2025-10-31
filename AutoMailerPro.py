#!/usr/bin/env python3
"""
Auto Mailer Pro v5.1
Author: Kyle Padilla
Company: Jones Insurance Advisors, Inc.
Contact: scooby_rizz@proton.me

Description:
    Automated script to generate personalized letters, envelopes,
    and mailing labels for owner-occupied properties (personal) or businesses (commercial).

Usage:
    python AutoMailerPro.py
    OR called from GUI with mode, file_path, content, subject_line, signature_name, signature_title, signature_image, and signature_email parameters.

Requirements:
    pandas, python-docx, fuzzywuzzy, python-Levenshtein
"""

__version__ = "5.1"
__author__ = "Kyle Padilla"
__company__ = "Jones Insurance Advisors, Inc."
__contact__ = "scooby_rizz@proton.me"

import os
import re
import shutil
import sqlite3
import sys
from datetime import datetime
from pathlib import Path
from typing import Dict, Iterable, List, Mapping, Optional
from typing import Dict, Iterable, List, Mapping

import pandas as pd

from docx import Document

from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt
from fuzzywuzzy import fuzz

import csv

def _get_first_nonempty(row, columns, default=""):
    """Return the first non-empty value found for the given columns in a row."""
    for column in columns:
        if column in row.index:
            value = row[column]
            if pd.notna(value):
                value_str = str(value).strip()
                if value_str and value_str.lower() != "nan":
                    return value_str
    return default
IGNORABLE_NAME_PREFIXES = {
    "mr",
    "mrs",
    "ms",
    "miss",
    "dr",
    "rev",
    "hon",
    "attn",
}

IGNORABLE_NAME_SUFFIXES = {
    "jr",
    "junior",
    "sr",
    "senior",
    "esq",
    "esquire",
    "jd",
    "md",
    "phd",
    "law",
}

ORDINAL_WORDS = {
    "first",
    "second",
    "third",
    "fourth",
    "fifth",
    "sixth",
    "seventh",
    "eighth",
    "ninth",
    "tenth",
    "1st",
    "2nd",
    "3rd",
    "4th",
    "5th",
    "6th",
    "7th",
    "8th",
    "9th",
    "10th",
}

ROMAN_NUMERAL_SUFFIXES = {
    "i",
    "ii",
    "iii",
    "iv",
    "v",
    "vi",
    "vii",
    "viii",
    "ix",
    "x",
}


def _normalize_name_token(token):
    """Return a simplified representation of a name token for comparisons."""
    return re.sub(r"[^a-z0-9]", "", token.lower())


def _strip_affixes(tokens):
    """Remove known prefixes/suffixes/ordinals from a list of name tokens."""
    filtered = []
    idx = 0
    while idx < len(tokens):
        token = tokens[idx]
        normalized = _normalize_name_token(token)
        if not normalized:
            idx += 1
            continue

        if normalized in IGNORABLE_NAME_PREFIXES:
            idx += 1
            continue

        if normalized in IGNORABLE_NAME_SUFFIXES or normalized in ROMAN_NUMERAL_SUFFIXES:
            idx += 1
            continue

        if normalized == "the" and idx + 1 < len(tokens):
            next_normalized = _normalize_name_token(tokens[idx + 1])
            if next_normalized in ORDINAL_WORDS or next_normalized in ROMAN_NUMERAL_SUFFIXES:
                idx += 2
                continue

        filtered.append(token)
        idx += 1

    return filtered


def _clean_name_tokens(name):
    if not name:
        return []
    tokens = str(name).replace(",", " ").split()
    return _strip_affixes(tokens)


def _format_given_names(tokens):
    """Format given-name tokens as ``First M.`` style."""
    if not tokens:
        return ""

    first_name = tokens[0].title()
    middle_initials = []
    for token in tokens[1:]:
        normalized = _normalize_name_token(token)
        if not normalized:
            continue
        middle_initials.append(f"{normalized[0].upper()}.")

    if middle_initials:
        return f"{first_name} {' '.join(middle_initials)}"
    return first_name



def _build_mailing_address(row):
    """Construct a mailing address string from any available components."""
    parts = []
    line_one = _get_first_nonempty(row, [
        "Mailing Address",
                "Mailing Address Line 1",
        "Mailing Address 1",
        "Address",
    ])
    if line_one:
        parts.append(line_one)

    line_two = _get_first_nonempty(row, [
        "Mailing Address 2",
        "Mailing Address Line 2",
        "Address Line 2",
    ])
    if line_two:
        parts.append(line_two)


    zip_code = _get_first_nonempty(row, [
        "Mailing Zip",
        "Mailing ZIP",
        "Mailing Zip Code",
        "Zip Code",
        "Zip",
    ])
    city_state_zip = _compose_city_state_zip(row, zip_code)
    if city_state_zip:
        parts.append(city_state_zip)

        return " | ".join(parts) if parts else ""

# === PATH CONFIGURATION ===



def get_resource_dir() -> Path:
    """Return the folder that contains bundled assets and data files."""
    if getattr(sys, "frozen", False) and hasattr(sys, "_MEIPASS"):
        return Path(sys._MEIPASS)
    return Path(__file__).resolve().parent


BASE_DIR = get_resource_dir()
DATA_DIR = BASE_DIR / "data"
ASSETS_DIR = BASE_DIR / "assets"
SIGNATURES_DIR = ASSETS_DIR / "signatures"


def get_user_data_dir(app_name: str = "AutoMailerPro") -> Path:
    """Return a user-writable folder for persistent application data."""

    if sys.platform.startswith("win"):
        base_dir = Path(os.environ.get("LOCALAPPDATA") or os.environ.get("APPDATA") or Path.home())
    elif sys.platform == "darwin":
        base_dir = Path.home() / "Library" / "Application Support"
    else:
        base_dir = Path(os.environ.get("XDG_DATA_HOME") or (Path.home() / ".local" / "share"))

    return base_dir / app_name


WRITABLE_DATA_DIR = get_user_data_dir()
OUTPUT_ROOT = WRITABLE_DATA_DIR / "output"
CAMPAIGN_DB_PATH = WRITABLE_DATA_DIR / "campaign_history.db"


CAMPAIGN_CONTACTS_COLUMNS = (
    "campaign_id",
    "mode",
    "sent_at",
    "name",
    "address",
    "zip",
    "sale_date",
    "sale_price",
    "email",
    "phone",
    "source",
)


CAMPAIGN_CONTACTS_TABLE_SQL = """
CREATE TABLE IF NOT EXISTS campaign_contacts (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    campaign_id TEXT NOT NULL,
    mode TEXT NOT NULL,
    sent_at TEXT NOT NULL,
    name TEXT NOT NULL,
    address TEXT NOT NULL,
    zip TEXT,
    sale_date TEXT,
    sale_price REAL,
    email TEXT,
    phone TEXT,
    source TEXT,
    created_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,
    UNIQUE (campaign_id, name, address)
)
"""


CAMPAIGN_CONTACTS_INSERT_SQL = """
INSERT INTO campaign_contacts (
    campaign_id,
    mode,
    sent_at,
    name,
    address,
    zip,
    sale_date,
    sale_price,
    email,
    phone,
    source
) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
ON CONFLICT(campaign_id, name, address) DO UPDATE SET
    mode=excluded.mode,
    sent_at=excluded.sent_at,
    zip=excluded.zip,
    sale_date=excluded.sale_date,
    sale_price=excluded.sale_price,
    email=excluded.email,
    phone=excluded.phone,
    source=excluded.source
"""


def _rebuild_campaign_contacts_table(connection):
    """Recreate the campaign_contacts table, preserving compatible legacy data."""

    cursor = connection.cursor()
    cursor.execute(
        "SELECT name FROM sqlite_master WHERE type='table' AND name='campaign_contacts'"
    )
    has_existing = cursor.fetchone() is not None

    if has_existing:
        cursor.execute("ALTER TABLE campaign_contacts RENAME TO campaign_contacts_legacy")

    cursor.execute(CAMPAIGN_CONTACTS_TABLE_SQL)

    if has_existing:
        cursor.execute("PRAGMA table_info('campaign_contacts_legacy')")
        legacy_columns = [row[1] for row in cursor.fetchall()]
        transferable_columns = [
            column for column in CAMPAIGN_CONTACTS_COLUMNS if column in legacy_columns
        ]
        if transferable_columns:
            column_list = ", ".join(transferable_columns)
            cursor.execute(
                f"INSERT OR IGNORE INTO campaign_contacts ({column_list}) "
                f"SELECT {column_list} FROM campaign_contacts_legacy"
            )
        cursor.execute("DROP TABLE IF EXISTS campaign_contacts_legacy")


def _ensure_campaign_history_schema(connection):
    """Ensure the campaign history database has the expected schema."""

    cursor = connection.cursor()
    cursor.execute(
        "SELECT sql FROM sqlite_master WHERE type='table' AND name='campaign_contacts'"
    )
    result = cursor.fetchone()

    if not result or not result[0]:
        cursor.execute(CAMPAIGN_CONTACTS_TABLE_SQL)
        return

    existing_sql = result[0].upper()
    if "UNIQUE (CAMPAIGN_ID, NAME, ADDRESS)" not in existing_sql:
        _rebuild_campaign_contacts_table(connection)
        return

    cursor.execute("PRAGMA table_info('campaign_contacts')")
    existing_columns = [row[1] for row in cursor.fetchall()]
    expected_columns = {"id", "created_at", *CAMPAIGN_CONTACTS_COLUMNS}

    if not expected_columns.issubset(set(existing_columns)):
        _rebuild_campaign_contacts_table(connection)


def _initialize_campaign_db(connection: sqlite3.Connection) -> None:
    """Ensure the SQLite database has the table for campaign contacts."""
def ensure_local_database() -> Path:
    """Create or hydrate the writable SQLite database in the user data directory."""

    WRITABLE_DATA_DIR.mkdir(parents=True, exist_ok=True)

    if not CAMPAIGN_DB_PATH.exists():
        packaged_db = DATA_DIR / "campaign_history.db"
        try:
            if packaged_db.exists():
                CAMPAIGN_DB_PATH.write_bytes(packaged_db.read_bytes())
            else:
                with sqlite3.connect(CAMPAIGN_DB_PATH) as connection:
                    _initialize_campaign_db(connection)
        except OSError as exc:
            raise RuntimeError(f"Unable to prepare local database: {exc}") from exc

    return CAMPAIGN_DB_PATH


CUSTOMERS_TABLE_SQL = """
CREATE TABLE IF NOT EXISTS customers (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    name TEXT NOT NULL,
    email TEXT,
    phone TEXT,
    premium REAL DEFAULT 0,
    home_price REAL DEFAULT 0,
    responded INTEGER DEFAULT 0,
    converted INTEGER DEFAULT 0,
    created_at TEXT DEFAULT CURRENT_TIMESTAMP,
    updated_at TEXT DEFAULT CURRENT_TIMESTAMP
)
"""


def _ensure_customers_table(connection: sqlite3.Connection) -> None:
    """Create the customers table if it is missing."""

    cursor = connection.cursor()
    cursor.execute(CUSTOMERS_TABLE_SQL)    
    _ensure_campaign_history_schema(connection)


def _append_campaign_records(
    records: Iterable[Mapping[str, object]],
    *,
    campaign_id: str,
    mode: str,
    sent_at: datetime,
) -> None:
    """Persist CRM rows into the consolidated SQLite database."""

    records = list(records)
    if not records:
        return

    WRITABLE_DATA_DIR.mkdir(parents=True, exist_ok=True)
    ensure_local_database()
    sent_at_iso = sent_at.isoformat(timespec="seconds")
    try:
        with sqlite3.connect(CAMPAIGN_DB_PATH) as connection:
            _initialize_campaign_db(connection)

            insert_sql = CAMPAIGN_CONTACTS_INSERT_SQL

            payload = []
            for record in records:
                payload.append(
                    (
                        campaign_id,
                        mode,
                        sent_at_iso,
                        str(record.get("Name", "")),
                        str(record.get("Address", "")),
                        str(record.get("Zip", "")) or None,
                        str(record.get("Sale Date", "")) or None,
                        float(record.get("Sale Price", 0.0) or 0.0),
                        str(record.get("Email", "")) or None,
                        str(record.get("Phone", "")) or None,
                        str(record.get("Source", "")) or None,
                    )
                )

            connection.executemany(insert_sql, payload)
            connection.commit()

            print(
                f"🗄️ Logged {len(payload)} contacts to campaign history database at {CAMPAIGN_DB_PATH}"
            )
    except sqlite3.Error as exc:
        print(f"⚠️ Failed to log campaign history: {exc}")

ZIP_LOOKUP_FILE = DATA_DIR / "zip_lookup.csv"
MASTER_CLIENT_LIST = DATA_DIR / "master_client_list.xlsx"
LOGO_PATH = ASSETS_DIR / "logo.png"
if not LOGO_PATH.exists():
    LOGO_PATH = BASE_DIR / "Logo.png"

YOUR_CO = "Jones Insurance Advisors, Inc"
YOUR_PHONE = "(772) 569-6802"
YOUR_ADDRESS = "3885 20th Street,\n Vero Beach, FL 32960"
YOUR_WEB = "www.jonesinsuranceadvisors.com"

# === ZIP TO CITY/STATE LOOKUP ===
zip_city_state = {}

def load_zip_lookup():
    global zip_city_state
    if not ZIP_LOOKUP_FILE.exists():
        print(f"❌ Missing ZIP lookup file: {ZIP_LOOKUP_FILE}")
        return
    df = pd.read_csv(ZIP_LOOKUP_FILE, dtype=str)
    for _, row in df.iterrows():
        raw_zip = row.get("zip")
        normalized_zip = _normalize_zip(raw_zip)
        if not normalized_zip:
            continue
        city = str(row.get("city", "")).strip()
        state = str(row.get("state", "")).strip()
        if not city or not state:
            continue
        city_state = f"{city.title()}, {state.upper()}"
        zip_city_state[normalized_zip] = city_state

def _normalize_zip(zip_code):
    """Return the 5-digit portion of a ZIP code string if available."""
    if zip_code is None:
        return ""

    digits_only = re.sub(r"\D", "", str(zip_code))
    if len(digits_only) >= 5:
        return digits_only[:5]
    if digits_only:
        return digits_only.zfill(5)
    return ""

def append_campaign_history(campaign_id, mode, crm_rows):
    """Append campaign contacts to the SQLite history database."""

    if not crm_rows:
        return

    ensure_local_database()

    sent_at = datetime.utcnow().isoformat(timespec="seconds")
    rows_to_insert = []
    for row in crm_rows:
        sale_price_raw = row.get("Sale Price", 0.0)
        try:
            sale_price = float(sale_price_raw or 0.0)
        except (TypeError, ValueError):
            sale_price = 0.0

        rows_to_insert.append(
            (
                campaign_id,
                mode,
                sent_at,
                row.get("Name", ""),
                row.get("Address", ""),
                row.get("Zip", ""),
                row.get("Sale Date", ""),
                sale_price,
                row.get("Email", ""),
                row.get("Phone", ""),
                row.get("Source", ""),
            )
        )

    try:
        with sqlite3.connect(CAMPAIGN_DB_PATH) as connection:
            _ensure_campaign_history_schema(connection)
            cursor = connection.cursor()
            cursor.executemany(CAMPAIGN_CONTACTS_INSERT_SQL, rows_to_insert)
            connection.commit()
        print(f"🗃️ Campaign history updated: {CAMPAIGN_DB_PATH}")
    except sqlite3.Error as error:
        print(f"⚠️ Failed to log campaign history: {error}")

def _prepare_customer_database(connection: sqlite3.Connection) -> None:
    """Ensure the campaign and customer tables exist for shared reporting."""

    _initialize_campaign_db(connection)
    _ensure_customers_table(connection)


def _to_float(value: object) -> float:
    """Safely convert arbitrary values to floats, returning 0.0 on failure."""

    if value in (None, ""):
        return 0.0
    try:
        return float(value)
    except (TypeError, ValueError):
        return 0.0


def list_customers() -> List[Dict[str, object]]:
    """Return all stored customers for display in the GUI."""

    WRITABLE_DATA_DIR.mkdir(parents=True, exist_ok=True)
    with sqlite3.connect(CAMPAIGN_DB_PATH) as connection:
        connection.row_factory = sqlite3.Row
        _prepare_customer_database(connection)
        cursor = connection.execute(
            """
            SELECT id, name, email, phone, premium, home_price, responded, converted,
                   created_at, updated_at
            FROM customers
            ORDER BY name COLLATE NOCASE
            """
        )
        return [dict(row) for row in cursor.fetchall()]


def save_customer(customer: Mapping[str, object]) -> int:
    """Insert or update a customer record in the database."""

    name = str(customer.get("name", "")).strip()
    if not name:
        raise ValueError("Customer name is required")

    email = str(customer.get("email", "")).strip() or None
    phone = str(customer.get("phone", "")).strip() or None
    premium = _to_float(customer.get("premium"))
    home_price = _to_float(customer.get("home_price"))
    responded = 1 if customer.get("responded") else 0
    converted = 1 if customer.get("converted") else 0
    customer_id = customer.get("id")

    WRITABLE_DATA_DIR.mkdir(parents=True, exist_ok=True)
    ensure_local_database()
    with sqlite3.connect(CAMPAIGN_DB_PATH) as connection:
        _prepare_customer_database(connection)
        cursor = connection.cursor()

        if customer_id:
            cursor.execute(
                """
                UPDATE customers
                SET name = ?,
                    email = ?,
                    phone = ?,
                    premium = ?,
                    home_price = ?,
                    responded = ?,
                    converted = ?,
                    updated_at = CURRENT_TIMESTAMP
                WHERE id = ?
                """,
                (name, email, phone, premium, home_price, responded, converted, customer_id),
            )
            if cursor.rowcount == 0:
                raise ValueError("Customer not found")
            connection.commit()
            return int(customer_id)

        cursor.execute(
            """
            INSERT INTO customers (name, email, phone, premium, home_price, responded, converted)
            VALUES (?, ?, ?, ?, ?, ?, ?)
            """,
            (name, email, phone, premium, home_price, responded, converted),
        )
        connection.commit()
        return int(cursor.lastrowid)


def _compute_group_metrics(customers: Iterable[Mapping[str, object]]) -> Dict[str, float]:
    records = list(customers)
    total_customers = len(records)
    responded_count = sum(1 for customer in records if customer.get("responded"))
    converted_count = sum(1 for customer in records if customer.get("converted"))

    response_rate = (responded_count / total_customers * 100) if total_customers else 0.0
    conversion_rate = (converted_count / total_customers * 100) if total_customers else 0.0

    home_prices = [
        _to_float(customer.get("home_price"))
        for customer in records
        if _to_float(customer.get("home_price")) > 0
    ]
    average_home_price = sum(home_prices) / len(home_prices) if home_prices else 0.0

    premiums = [
        _to_float(customer.get("premium"))
        for customer in records
        if _to_float(customer.get("premium")) > 0
    ]
    average_premium = sum(premiums) / len(premiums) if premiums else 0.0
    total_premium = sum(premiums)
    responded_premium = sum(
        _to_float(customer.get("premium")) for customer in records if customer.get("responded")
    )
    converted_premium = sum(
        _to_float(customer.get("premium")) for customer in records if customer.get("converted")
    )
    prospect_premium = sum(
        _to_float(customer.get("premium"))
        for customer in records
        if not customer.get("responded")
    )
    return {
        "total_customers": total_customers,
        "responded_count": responded_count,
        "converted_count": converted_count,
        "response_rate": response_rate,
        "conversion_rate": conversion_rate,
        "average_home_price": average_home_price,
        "average_premium": average_premium,
                "total_premium": total_premium,
        "responded_premium": responded_premium,
        "converted_premium": converted_premium,
        "prospect_premium": prospect_premium,
    }


def get_customer_metrics(filters: Optional[Mapping[str, bool]] = None) -> Dict[str, object]:
    """Calculate aggregate metrics used by the reporting window."""

    customers = list_customers()
    overall_metrics = _compute_group_metrics(customers)

    filters = filters or {}
    include_prospects = filters.get("include_prospects", True)
    include_responded = filters.get("include_responded", True)
    include_converted = filters.get("include_converted", True)

    status_groups = {
        "prospect": [],
        "responded": [],
        "converted": [],
    }

    for customer in customers:
        if customer.get("converted"):
            status_groups["converted"].append(customer)
        elif customer.get("responded"):
            status_groups["responded"].append(customer)
        else:
            status_groups["prospect"].append(customer)

    allowed_statuses = set()
    if include_prospects:
        allowed_statuses.add("prospect")
    if include_responded:
        allowed_statuses.add("responded")
    if include_converted:
        allowed_statuses.add("converted")

    filtered_customers: List[Mapping[str, object]] = []
    for status in allowed_statuses:
        filtered_customers.extend(status_groups[status])

    filtered_metrics = _compute_group_metrics(filtered_customers)
    status_breakdown = {
        "prospects": _compute_group_metrics(status_groups["prospect"]),
        "responded": _compute_group_metrics(status_groups["responded"]),
        "converted": _compute_group_metrics(status_groups["converted"]),
    }

    return {
        **overall_metrics,
        "filters": {
            "include_prospects": include_prospects,
            "include_responded": include_responded,
            "include_converted": include_converted,
        },
        "filtered_totals": filtered_metrics,
        "status_breakdown": status_breakdown,
    }





def _compose_city_state_zip(row, zip_code):
    """Create a display string for city/state/ZIP using row data before falling back."""
    city = _get_first_nonempty(row, ["Mailing City", "City"]).strip()
    state = _get_first_nonempty(row, ["Mailing State", "State"]).strip()

    city_formatted = city.title() if city else ""
    state_formatted = state.upper() if state else ""

    location = ""
    if city_formatted and state_formatted:
        location = f"{city_formatted}, {state_formatted}"
    elif city_formatted:
        location = city_formatted
    elif state_formatted:
        location = state_formatted

    display_zip = ""
    if zip_code is not None:
        raw_zip = str(zip_code).strip()
        if raw_zip and raw_zip.lower() != "nan":
            normalized_zip = _normalize_zip(zip_code)
            display_zip = normalized_zip or raw_zip

    if location:
        if display_zip:
            return f"{location} {display_zip}".strip()
        return location

    fallback = zip_to_city_state(zip_code)
    if fallback:
        return fallback
    return display_zip


def zip_to_city_state(zip_code):
    normalized_zip = _normalize_zip(zip_code)
    city_state = zip_city_state.get(normalized_zip, "Indian River County, FL") if normalized_zip else "Indian River County, FL"
    display_zip = normalized_zip or str(zip_code).strip()
    if not display_zip or display_zip.lower() in {"nan", ""}:
        return city_state
    return f"{city_state} {display_zip}".strip()

# === LOAD CLIENT LIST FOR SCRUBBING ===
def load_client_list():
    if not MASTER_CLIENT_LIST.exists():
        print(f"❌ Master client list not found: {MASTER_CLIENT_LIST}")
        return []
    try:
        df = pd.read_excel(MASTER_CLIENT_LIST)
        return df[['Name', 'Mailing Address']].dropna().to_dict('records')
    except Exception as e:
        print(f"⚠️ Failed to load master client list: {e}")
        return []

# === CHECK IF RECORD IS IN CLIENT LIST ===
def is_existing_client(name, mailing_address, client_list):
    if not client_list:
        return False
    name = str(name).lower().strip()
    mailing_address = str(mailing_address).lower().strip()
    for client in client_list:
        client_name = str(client.get('Name', '')).lower().strip()
        client_address = str(client.get('Mailing Address', '')).lower().strip()
        if (fuzz.partial_ratio(name, client_name) > 85 and
            fuzz.partial_ratio(mailing_address, client_address) > 85):
            return True
    return False

# === CLEAN NAME ===
def clean_name(row, mode):
    if mode == "personal":
        raw_name = _get_first_nonempty(row, ['Owner Name', 'Owner'])
        name_parts = [part.strip() for part in raw_name.split('||')]
        last_names = []
        first_names = []
        for part in name_parts:
            name_no_suffix = part.split('(')[0].strip()
            words = _clean_name_tokens(name_no_suffix)
            if len(words) >= 2:
                last_names.append(words[0].title())
                first_name = _format_given_names(words[1:])
                first_names.append(first_name)
            elif len(words) == 1:
                last_names.append("")
                ffirst_names.append(_format_given_names([words[0]]))
            else:
                last_names.append("")
                first_names.append("")
        unique_last_names = set([ln for ln in last_names if ln])
        if len(unique_last_names) == 1:
            last_name = unique_last_names.pop()
            combined_first_names = " & ".join([fn for fn in first_names if fn])
            full_name = f"{combined_first_names} {last_name}".strip()
        else:
            full_name = " & ".join(
                [f"{fn} {ln}".strip() for fn, ln in zip(first_names, last_names) if fn or ln]
            )
        return full_name or "Valued Customer"
    else:  # commercial
        first_name = str(row.get('Executive First Name', '')).strip()
        last_name = str(row.get('Executive Last Name', '')).strip()
        first_tokens = _clean_name_tokens(first_name)
        last_tokens = _clean_name_tokens(last_name)
        cleaned_first_name = _format_given_names(first_tokens)
        cleaned_last_name = " ".join(token.title() for token in last_tokens)
        if cleaned_first_name and cleaned_last_name:
            return f"{cleaned_first_name} {cleaned_last_name}".strip()
        legal_name = str(row.get('Legal Name', '')).strip()
        company_name = str(row.get('Company Name', '')).strip()
        return legal_name.title() or company_name.title() or "Valued Business"

def _has_minimum_name_parts(name, min_parts=2):
    """Return True if the cleaned name contains at least ``min_parts`` distinct words."""
    if not name:
        return False

    word_count = sum(
        1 for word in name.replace('&', ' ').split() if any(char.isalpha() for char in word)
    )
    return word_count >= min_parts

# === FILTERS ===
def is_owner_occupied(property_address, mailing_address):
    try:
        prop_addr = str(property_address).lower().strip()
        mailing_parts = str(mailing_address).split('|')
        for part in mailing_parts:
            if fuzz.partial_ratio(prop_addr, part.lower().strip()) > 85:
                return True
        return False
    except:
        return False

# === BUSINESS TYPE VALIDATION ===
def is_valid_business(business_type):
    """Return True if the business type represents a qualified commercial lead."""
    if business_type is None:
        return False

    business_type = str(business_type).strip().lower()
    if not business_type:
        return False

    disqualifying_keywords = {
        "church",
        "religious",
        "synagogue",
        "temple",
        "mosque",
        "school",
        "college",
        "university",
        "government",
        "county",
        "city",
        "state",
        "federal",
        "municipal",
        "public",
        "utility",
        "hoa",
        "homeowners association",
        "condominium",
        "condo",
        "apartments",
        "apartment",
        "association",
        "non-profit",
        "nonprofit",
        "charity",
        "vacant",
        "land",
        "empty lot",
        "lot",
    }

    for keyword in disqualifying_keywords:
        if keyword in business_type:
            return False

    return True

# === ADD LETTER TO DOC ===
def add_letter_to_doc(doc, name, address, zip_code, sale_date, sale_price, content, mode, subject_line, signature_name, signature_title, signature_image, signature_email):
    today = datetime.now().strftime('%B %d, %Y')

    def add_compact_paragraph(text="", bold=False, space_before=0, space_after=2):
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(space_before)
        para.paragraph_format.space_after = Pt(space_after)
        run = para.add_run(text)
        run.bold = bold
        return para

    for _ in range(4):
        doc.add_paragraph()

    add_compact_paragraph(today, space_after=48)
    greeting = f"Dear {name},"
    add_compact_paragraph(greeting, space_after=24)

    add_compact_paragraph(subject_line, bold=True, space_after=10)

    city_state = zip_to_city_state(zip_code)
    county_name = "Indian River" if "County" in city_state else city_state.split("County")[0].strip()
    personalized_content = content.replace("[Name]", name).replace("[County]", county_name)

    doc.add_paragraph(personalized_content)

    signature_image_path = os.fspath(signature_image) if signature_image else None
    if signature_image_path and os.path.exists(signature_image_path):
        doc.add_picture(signature_image_path, width=Inches(1.5), height=Inches(0.5))
    else:
        print(f"❌ Signature image not found: {signature_image}")

    doc.add_paragraph(
        f"{signature_name}\n{signature_title}\n{signature_email}\n{YOUR_PHONE}\n{YOUR_WEB}"
    )

    doc.add_page_break()

# === ADD ENVELOPE TO DOC ===
def add_envelope_to_doc(doc, name, address, location_line, signature_name):
    section = doc.add_section()
    section.page_width = Inches(9.5)
    section.page_height = Inches(4.125)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)

    return_address = f"{signature_name}\n{YOUR_ADDRESS}"
    sender = doc.add_paragraph(return_address)
    sender.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    for run in sender.runs:
        run.font.size = Pt(10)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(40)

    recipient = doc.add_paragraph()
    recipient.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    name_run = recipient.add_run(f"{name}\n")
    name_run.bold = True
    name_run.font.size = Pt(14)
    addr_line = f"{address}\n{location_line}" if location_line else address
    addr_run = recipient.add_run(addr_line)
    addr_run.font.size = Pt(14)

    doc.add_page_break()

# === CREATE LABELS DOC ===
def create_labels(label_data, labels_file):
    labels_file = Path(labels_file)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.19)
    section.right_margin = Inches(0.19)

    labels_per_row = 3
    label_width = Inches(2.63)
    label_height = Inches(1.0)
    num_rows = (len(label_data) + labels_per_row - 1) // labels_per_row

    table = doc.add_table(rows=num_rows, cols=labels_per_row)
    table.autofit = False
    table.allow_autofit = False

    for col in table.columns:
        col.width = label_width
    for row in table.rows:
        row.height = label_height
        row.height_rule = 2

    idx = 0
    for row in table.rows:
        for cell in row.cells:
            if idx < len(label_data):
                lines = label_data[idx].split("\n")
                para = cell.paragraphs[0]
                para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                if lines:
                    name_run = para.add_run(lines[0] + "\n")
                    name_run.bold = True
                    name_run.font.size = Pt(10.5)
                for line in lines[1:]:
                    addr_run = para.add_run(line + "\n")
                    addr_run.font.size = Pt(10.5)
                idx += 1
            else:
                cell.text = ""

    doc.save(str(labels_file))
    print(f"✅ Mailing labels saved to: {labels_file}")

# === MAIN ===
def main(
    mode="personal",
    file_path=DATA_DIR / "sales_data.xlsx",
    content=None,
    subject_line="",
    signature_name="Brian Jones",
    signature_title="Vice President",
    signature_image=SIGNATURES_DIR / "signature_brian.png",
    signature_email="Brian@jonesia.com",
):
    if mode not in ["personal", "commercial"]:
        raise ValueError("Mode must be 'personal' or 'commercial'")
    if not subject_line:
        if mode == "personal":
            subject_line = "Homeowners Insurance Rates Are Finally on the Decline – Don’t Miss Out!"
        else:
            subject_line = "Protect Your Business with Tailored Insurance Solutions!"

    if content is None:
        if mode == "personal":
            content = (
                "For the first time in years, homeowners rates are coming down — and the savings could be significant.\n\n"
                "Recent legislative changes have boosted competition in Florida’s property insurance market, "
                "and many Indian River County homeowners are already benefiting.\n\n"
                "Jones Insurance Advisors is a two-generation, family-owned independent agency located right here in Vero Beach. "
                "Our team of dedicated agents possess extensive knowledge of the intricacies of the local insurance market, "
                "and are excited to assist you in finding the most comprehensive and competitively priced insurance solutions.\n\n"
                "Call us today for a free, no-obligation quote, or visit our website below and complete a quote request, "
                "and one of our dedicated agents will reach out to you!\n\n"
                "We look forward to earning your business and providing you the personal, dedicated service you have come to "
                "expect by doing business locally.\n\n"
                "Warm Regards,"
            )
        else:
            content = (
                "Protecting your business is our priority at Jones Insurance Advisors.\n\n"
                "As an Indian River County business, you need insurance solutions tailored to your unique needs. "
                "Our experienced team specializes in crafting comprehensive coverage plans for businesses like yours, "
                "ensuring protection against risks while keeping costs competitive.\n\n"
                "Jones Insurance Advisors, a family-owned agency in Vero Beach, is here to help. "
                "Contact us for a free consultation to discuss how we can safeguard your business.\n\n"
                "We look forward to partnering with you!\n\n"
                "Best Regards,"
            )



    letters_doc = Document()
    envelopes_doc = Document()

    load_zip_lookup()
    client_list = load_client_list()

    file_path = Path(file_path)
    if not file_path.exists():
        raise FileNotFoundError(f"Excel file not found: {file_path}")

    signature_image = Path(signature_image)
    if not signature_image.exists() and not signature_image.is_absolute():
        candidate = SIGNATURES_DIR / signature_image.name
        if candidate.exists():
            signature_image = candidate

    try:
        df = pd.read_excel(file_path)
    except Exception as e:
        raise Exception(f"Failed to read Excel file: {e}")
    sale_date_range_label = None
    if 'Sale Date' in df.columns:
        sale_dates = pd.to_datetime(df['Sale Date'], errors='coerce')
        sale_dates = sale_dates.dropna()
        if not sale_dates.empty:
            oldest_sale = sale_dates.min()
            newest_sale = sale_dates.max()
            sale_date_range_label = f"{oldest_sale.strftime('%m%d%y')}-{newest_sale.strftime('%m%d%y')}"
    run_started_at = datetime.now()
    timestamp = run_started_at.strftime("%m%d%y_%H%M%S")
    OUTPUT_ROOT.mkdir(parents=True, exist_ok=True)
    if sale_date_range_label:
        folder_name = f"{sale_date_range_label}_{mode.capitalize()}_Mailing_Campaign"
    else:
        
        folder_name = f"{timestamp}_{mode.capitalize()}_Mailing_Campaign"

    OUTPUT_DIR = OUTPUT_ROOT / folder_name
    LETTERS_FILE = OUTPUT_DIR / "all_letters.docx"
    ENVELOPES_FILE = OUTPUT_DIR / "all_envelopes.docx"
    LABELS_FILE = OUTPUT_DIR / "mailing_labels.docx"
    CRM_EXPORT_FILE = OUTPUT_DIR / f"crm_{mode}_occupied.csv"

    created_output_dir = not OUTPUT_DIR.exists()
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    if created_output_dir:
        print(f"📁 Created output folder: {OUTPUT_DIR}")

    labels = []
    crm_rows = []
    is_new_format = 'Executive First Name' in df.columns and 'Executive Last Name' in df.columns

    for _, row in df.iterrows():
        try:
            property_address = _get_first_nonempty(row, ['Address', 'Situs'])
            mailing_address_value = _build_mailing_address(row)
            if mode == "personal":
                name_key = 'Owner Name'
                filter_check = is_owner_occupied(property_address, mailing_address_value)
                filter_desc = "non-owner-occupied"
            else:  # commercial
                name_key = 'Executive First Name' if is_new_format else 'Business Name'
                filter_check = is_valid_business(row.get('Business Type', '')) if not is_new_format else True
                filter_desc = "invalid business type"

            name = clean_name(row, mode)
            if not name:
                print(f"⏭️ Skipping row with missing name")
                continue
            if not _has_minimum_name_parts(name):
                print(f"⏭️ Skipping insufficient name parts: {name}")
                continue

            address = _get_first_nonempty(row, ['Address', 'Situs']).title().strip()
            zip_code = _get_first_nonempty(row, ['Site Zip Code', 'Property Zip', 'Zip Code', 'Zip'])
            mailing_address = mailing_address_value if mode == "personal" else _get_first_nonempty(row, ['Address'])
            location_line = _compose_city_state_zip(row, zip_code)
            sale_date_raw = _get_first_nonempty(row, ['Sale Date']) if not is_new_format else "Unknown"
            sale_price_str = _get_first_nonempty(row, ['Sale Price']) if not is_new_format else "0.0"
            sale_price_str = sale_price_str.replace('$', '').replace(',', '') if sale_price_str else ''

            if is_existing_client(name, mailing_address, client_list):
                print(f"⏭️ Skipping existing client: {name}")
                continue

            if not filter_check and not is_new_format:
                print(f"⏭️ Skipping {filter_desc}: {name}")
                continue

            try:
                sale_price = float(sale_price_str) if sale_price_str else 0.0
            except ValueError:
                sale_price = 0.0
            try:
                sale_date = datetime.strptime(sale_date_raw, '%m/%d/%Y').strftime('%B %d, %Y') if sale_date_raw else "Unknown"
            except ValueError:
                sale_date = "Unknown"

            add_letter_to_doc(letters_doc, name, address, zip_code, sale_date, sale_price, content, mode, subject_line, signature_name, signature_title, signature_image, signature_email)
            add_envelope_to_doc(envelopes_doc, name, address, location_line, signature_name)

            label_text = f"{name}\n{address}\n{location_line}" if location_line else f"{name}\n{address}"
            labels.append(label_text)

            crm_rows.append({
                'Name': name,
                'Address': address,
                'Zip': zip_code,
                'Sale Date': sale_date,
                'Sale Price': sale_price,
                'Email': '',
                'Phone': '',
                'Source': f"{mode.capitalize()} Anniversary Mailer-Sept-Oct"
            })

            print(f"✅ Processed: {name}")

        except Exception as e:
            print(f"⚠️ Skipped row due to error: {e}")

    if labels:
        create_labels(labels, LABELS_FILE)

    if crm_rows:
        keys = crm_rows[0].keys()
        with open(CRM_EXPORT_FILE, 'w', newline='', encoding='utf-8') as f:
            dict_writer = csv.DictWriter(f, keys)
            dict_writer.writeheader()
            dict_writer.writerows(crm_rows)
        print(f"📥 CRM-ready CSV saved to: {CRM_EXPORT_FILE}")
        append_campaign_history(folder_name, mode, crm_rows)
        _append_campaign_records(
            crm_rows,
            campaign_id=OUTPUT_DIR.name,
            mode=mode,
            sent_at=run_started_at,
        )
    letters_doc.save(str(LETTERS_FILE))
    envelopes_doc.save(str(ENVELOPES_FILE))
    print(f"📄 All letters saved to: {LETTERS_FILE}")
    print(f"✉️ All envelopes saved to: {ENVELOPES_FILE}")

def print_logo():
    logo = r"""
                      __/___             
             _____/______|           
     _______/_____\_______\_____     
     \              < < <       |    
      ~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~  
       Auto Mailer Pro v5.1
       Author: Kyle Padilla
       Last Updated: 08/14/2025
       
           Version: 05.01
    """
    print(logo)

if __name__ == "__main__":
    print_logo()
    main()