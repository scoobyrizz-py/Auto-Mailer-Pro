import os
import pandas as pd
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from fuzzywuzzy import fuzz
import csv

# === CONFIG ===
EXCEL_FILE = "sales_data.xlsx"
ZIP_LOOKUP_FILE = "zip_lookup.csv"
LOGO_PATH = "logo.png"

YOUR_NAME = "Brain Jones"
YOUR_TITLE = "Vice President"
YOUR_CO = "Jones Insurance Advisors, Inc"
YOUR_PHONE = "(772) 569-6802"
YOUR_EMAIL = "Brian@jonesia.com"
YOUR_ADDRESS = "3885 20th Street,\n Vero Beach, FL 32960"
YOUR_WEB = "www.jonesinsuranceadvisors.com"
YOUR_RETURN_ADDRESS = f"{YOUR_NAME}\n{YOUR_ADDRESS}"

LETTERS_FILE = "all_letters.docx"
ENVELOPES_FILE = "all_envelopes.docx"
LABELS_FILE = "mailing_labels.docx"
CRM_EXPORT_FILE = "crm_owner_occupied.csv"

# === DOCUMENTS ===
letters_doc = Document()
envelopes_doc = Document()

# === ZIP TO CITY/STATE LOOKUP ===
zip_city_state = {}

def load_zip_lookup():
    global zip_city_state
    if not os.path.exists(ZIP_LOOKUP_FILE):
        print(f"❌ Missing ZIP lookup file: {ZIP_LOOKUP_FILE}")
        return

    df = pd.read_csv(ZIP_LOOKUP_FILE, dtype=str)
    for _, row in df.iterrows():
        zip_code = row['zip'].zfill(5)
        city_state = f"{row['city'].title()}, {row['state'].upper()}"
        zip_city_state[zip_code] = city_state

def zip_to_city_state(zip_code):
    zip_code = zip_code.zfill(5)  # Ensure 5-digit format
    city_state = zip_city_state.get(zip_code, "Indian River County, FL")
    return f"{city_state} {zip_code}"  # Append ZIP to the end

# === CLEAN OWNER NAME ===
def clean_owner_name(raw_name):
    name_parts = [part.strip() for part in raw_name.split('||')]
    cleaned_names = []
    for part in name_parts:
        # Remove suffix in parentheses like (LE)
        name_no_suffix = part.split('(')[0].strip()
        words = name_no_suffix.split()
        if len(words) >= 2:
            last_name = words[0]
            first_name = words[1]
            full_name = f"{first_name} {last_name}"
        elif len(words) == 1:
            full_name = words[0]
        else:
            full_name = ""
        full_name = full_name.title()
        cleaned_names.append(full_name)
    return ' & '.join(cleaned_names)

# === OWNER-OCCUPIED CHECK ===
def is_owner_occupied(property_address, mailing_address):
    try:
        prop_addr = property_address.lower().strip()
        mailing_parts = mailing_address.split('|')
        for part in mailing_parts:
            if fuzz.partial_ratio(prop_addr, part.lower().strip()) > 85:
                return True
        return False
    except:
        return False

# === ADD LETTER TO DOC ===
def add_letter_to_doc(doc, owner_name, address, zip_code, sale_date, sale_price):
    today = datetime.now().strftime('%B %d, %Y')

    # Add top padding for letterhead
    for _ in range(2):
        doc.add_paragraph()

    # Date at the top
    doc.add_paragraph(today)

    # Blank line
    doc.add_paragraph()

    # Determine county name from city/state lookup (fallback to generic)
    city_state = zip_to_city_state(zip_code)
    county_name = "your county"
    if "County" in city_state:
        county_name = city_state.split("County")[0].strip()
    elif "Vero Beach" in city_state:
        county_name = "Indian River"
    elif "Sebastian" in city_state:
        county_name = "Indian River"

    # Greeting
    doc.add_paragraph(f"Dear {owner_name},", style='Normal')
    doc.add_paragraph()

    # Letter content
    body = (
    f"IMPORTANT INSURANCE ALERT:\n\n"
    f"Homeowners insurance rates are finally stabilizing — for the first time in years. "
    f"And just this year, a highly competitive insurance carrier has expanded into {county_name}\n\n"
    "Why this matters right now:\n\n"
    f"• Recent legislative changes have opened the door for more competition in {county_name} County\n"
    "• Homeowners are saving thousands of dollars because of these changes\n"
    f"• Based on your property details, you may qualify for significant savings\n\n"
    "Jones Insurance Advisors is a local, family-owned agency right here in Vero Beach. "
    "We can provide you with a no-pressure, free quote — it’s fast, easy, and tailored to you.\n\n"
    )

    doc.add_paragraph(body)

    # Closing
    doc.add_paragraph("Warm regards,\n")

    # Signature image
    if os.path.exists("signature_brian.png"):
        doc.add_picture("signature_brian.png", width=Inches(1.5), height=Inches(0.5))

    # Signature block
    doc.add_paragraph(
        f"{YOUR_NAME}\n{YOUR_TITLE}\n{YOUR_CO}\n{YOUR_PHONE}\n{YOUR_EMAIL}\n{YOUR_ADDRESS}\n{YOUR_WEB}"
    )

    doc.add_page_break()



# === ADD ENVELOPE TO DOC ===
def add_envelope_to_doc(doc, owner_name, address, zip_code):
    section = doc.add_section(0)
    section.page_width = Inches(9.5)
    section.page_height = Inches(4.125)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)

    sender = doc.add_paragraph(YOUR_RETURN_ADDRESS)
    sender.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    sender.runs[0].font.size = Pt(10)

    doc.add_paragraph("\n\n\n\n\n")

    recipient = doc.add_paragraph()
    recipient.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = recipient.add_run(f"{owner_name}\n{address}\n{zip_to_city_state(zip_code)}")
    run.font.size = Pt(14)

# === CREATE LABELS DOC ===
def create_labels(label_data):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.19)
    section.right_margin = Inches(0.19)

    table = doc.add_table(rows=0, cols=3)
    table.autofit = False
    col_width = Inches(2.63)
    for col in table.columns:
        col.width = col_width

    for i, data in enumerate(label_data):
        if i % 3 == 0:
            row_cells = table.add_row().cells
        cell = row_cells[i % 3]
        para = cell.paragraphs[0]
        para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        run = para.add_run(data)
        run.font.size = Pt(10)

    doc.save(LABELS_FILE)
    print(f"✅ Mailing labels saved to: {LABELS_FILE}")

# === MAIN ===
def main():
    load_zip_lookup()

    if not os.path.exists(EXCEL_FILE):
        print(f"❌ File not found: {EXCEL_FILE}")
        return

    try:
        df = pd.read_excel(EXCEL_FILE)
    except Exception as e:
        print(f"⚠️ Failed to read Excel file: {e}")
        return

    labels = []
    crm_rows = []

    for _, row in df.iterrows():
        try:
            raw_owner = str(row['Owner Name'])
            address = str(row['Address']).title().strip()
            zip_code = str(row['Site Zip Code']).strip()
            mailing_address = str(row['Mailing Address']).strip()
            sale_date_raw = str(row['Sale Date']).strip()
            sale_price_str = str(row['Sale Price']).replace('$', '').replace(',', '').strip()
            sale_price = float(sale_price_str)

            if not is_owner_occupied(address, mailing_address):
                print(f"⏭️ Skipping non-owner-occupied: {raw_owner}")
                continue

            owner_name = clean_owner_name(raw_owner)
            sale_date = datetime.strptime(sale_date_raw, '%m/%d/%Y').strftime('%B %d, %Y')

            add_letter_to_doc(letters_doc, owner_name, address, zip_code, sale_date, sale_price)
            add_envelope_to_doc(envelopes_doc, owner_name, address, zip_code)

            label_text = f"{owner_name}\n{address}\n{zip_to_city_state(zip_code)}"
            labels.append(label_text)

            crm_rows.append({
                'Name': owner_name,
                'Address': address,
                'Zip': zip_code,
                'Sale Date': sale_date,
                'Sale Price': sale_price,
                'Email': '',
                'Phone': '',
                'Source': 'Homeowner Anniversary Mailer'
            })

            print(f"✅ Processed: {owner_name}")

        except Exception as e:
            print(f"⚠️ Skipped row due to error: {e}")

    letters_doc.save(LETTERS_FILE)
    envelopes_doc.save(ENVELOPES_FILE)
    print(f"📄 All letters saved to: {LETTERS_FILE}")
    print(f"✉️ All envelopes saved to: {ENVELOPES_FILE}")

    if labels:
        create_labels(labels)

    if crm_rows:
        keys = crm_rows[0].keys()
        with open(CRM_EXPORT_FILE, 'w', newline='', encoding='utf-8') as f:
            dict_writer = csv.DictWriter(f, keys)
            dict_writer.writeheader()
            dict_writer.writerows(crm_rows)
        print(f"📥 CRM-ready CSV saved to: {CRM_EXPORT_FILE}")

if __name__ == "__main__":
    main()
