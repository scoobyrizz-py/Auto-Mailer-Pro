#!/usr/bin/env python3
"""
InsuranceMailer v1.0
Author: Kyle Padilla
Company: Jones Insurance Advisors, Inc.
Contact: scooby_rizz@proton.me

Description:
    Automated script to generate personalized letters, envelopes,
    and mailing labels for owner-occupied properties using sales data.

Usage:
    python InsuranceMailer_v1.0.py

Requirements:
    pandas, python-docx, fuzzywuzzy, python-Levenshtein
"""

__version__ = "1.0"
__author__ = "Kyle Padilla"
__company__ = "Jones Insurance Advisors, Inc."
__contact__ = "scooby_rizz@proton.me"

def main():
    print(f"InsuranceMailer v{__version__} by {__author__} - {__company__}")
    # Your main program code here...

if __name__ == "__main__":
    main()

import os
import pandas as pd
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from fuzzywuzzy import fuzz
import csv




# === CONFIG ===
EXCEL_FILE = "sales_data.xlsx"
ZIP_LOOKUP_FILE = "zip_lookup.csv"
LOGO_PATH = "logo.png"

YOUR_NAME = "Brian Jones"
YOUR_TITLE = "Vice President"
YOUR_CO = "Jones Insurance Advisors, Inc"
YOUR_PHONE = "(772) 569-6802"
YOUR_EMAIL = "Brian@jonesia.com"
YOUR_ADDRESS = "3885 20th Street,\n Vero Beach, FL 32960"
YOUR_WEB = "www.jonesinsuranceadvisors.com"
YOUR_RETURN_ADDRESS = f"{YOUR_NAME}\n{YOUR_ADDRESS}"

# Create output directory with timestamp
OUTPUT_DIR = "090124_101524_Mailing Campaign"
LETTERS_FILE = os.path.join(OUTPUT_DIR, "all_letters.docx")
ENVELOPES_FILE = os.path.join(OUTPUT_DIR, "all_envelopes.docx")
LABELS_FILE = os.path.join(OUTPUT_DIR, "mailing_labels.docx")
CRM_EXPORT_FILE = os.path.join(OUTPUT_DIR, "crm_owner_occupied.csv")

# Create output directory if it doesn't exist
if not os.path.exists(OUTPUT_DIR):
    os.makedirs(OUTPUT_DIR)
    print(f"📁 Created output folder: {OUTPUT_DIR}")
    
# === DOCUMENTS ===
letters_doc = Document()
envelopes_doc = Document()

# === ZIP TO CITY/STATE LOOKUP ===
zip_city_state = {}

def load_zip_lookup():
    global zip_city_state
    if not os.path.exists(ZIP_LOOKUP_FILE):
        print(f"❌ Missing ZIP lookup file: {ZIP_LOOKUP_FILE}")
        return

    df = pd.read_csv(ZIP_LOOKUP_FILE, dtype=str)
    for _, row in df.iterrows():
        zip_code = row['zip'].zfill(5)
        city_state = f"{row['city'].title()}, {row['state'].upper()}"
        zip_city_state[zip_code] = city_state

def zip_to_city_state(zip_code):
    zip_code = zip_code.zfill(5)  # Ensure 5-digit format
    city_state = zip_city_state.get(zip_code, "Indian River County, FL")
    return f"{city_state} {zip_code}"  # Append ZIP to the end

# === CLEAN OWNER NAME ===
def clean_owner_name(raw_name):
    name_parts = [part.strip() for part in raw_name.split('||')]
    last_names = []
    first_names = []

    for part in name_parts:
        # Remove suffix in parentheses like (LE)
        name_no_suffix = part.split('(')[0].strip()
        words = name_no_suffix.split()
        if len(words) >= 2:
            last_names.append(words[0].title())
            first_names.append(words[1].title())
        elif len(words) == 1:
            # If only one word, treat it as first name (no last name)
            last_names.append("")
            first_names.append(words[0].title())
        else:
            last_names.append("")
            first_names.append("")

    # Check if all last names are the same and non-empty
    unique_last_names = set([ln for ln in last_names if ln])

    if len(unique_last_names) == 1:
        # Same last name for all, join first names & append last name once
        last_name = unique_last_names.pop()
        combined_first_names = " & ".join(first_names)
        full_name = f"{combined_first_names} {last_name}"
    else:
        # Different last names, join full names normally
        full_name = " & ".join(
            [f"{fn} {ln}".strip() for fn, ln in zip(first_names, last_names)]
        )

    return full_name


# === OWNER-OCCUPIED CHECK ===
def is_owner_occupied(property_address, mailing_address):
    try:
        prop_addr = property_address.lower().strip()
        mailing_parts = mailing_address.split('|')
        for part in mailing_parts:
            if fuzz.partial_ratio(prop_addr, part.lower().strip()) > 85:
                return True
        return False
    except:
        return False
# === Yellow pages API
import requests

# Replace with your RapidAPI credentials
API_KEY = "b6ff07f33bmsh5d8a96a3d9c7988p1d77b5jsn7ea8f0880f7f"

def lookup_owner_info(owner_name, city, state):
    url = "https://yellow-pages3.p.rapidapi.com/businesses/search"
    
    querystring = {
        "query": owner_name,
        "location": f"{city}, {state}",
        "page": "1",
        "size": "1"  # just the top match
    }

    headers = {
        "X-RapidAPI-Key": API_KEY,
        "X-RapidAPI-Host": "yellow-pages3.p.rapidapi.com"
    }

    try:
        response = requests.get(url, headers=headers, params=querystring)
        data = response.json()
        
        if "data" in data and len(data["data"]) > 0:
            business = data["data"][0]
            phone = business.get("phone", "")
            email = business.get("email", "")
            return phone, email
        else:
            return "", ""
    except Exception as e:
        print(f"API lookup error for {owner_name}: {e}")
        return "", ""

    # "X-RapidAPI-Key": "",


# === ADD LETTER TO DOC ===
def add_letter_to_doc(doc, owner_name, address, zip_code, sale_date, sale_price):
    today = datetime.now().strftime('%B %d, %Y')

    def add_compact_paragraph(text="", bold=False, space_before=0, space_after=2):
        """Add a paragraph with tight spacing."""
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(space_before)
        para.paragraph_format.space_after = Pt(space_after)
        run = para.add_run(text)
        run.bold = bold
        return para

    # Add top padding for letterhead (normal spacing here)
    for _ in range(4):
        doc.add_paragraph()

    # Date at the top (tight spacing)
    add_compact_paragraph(today, space_after=48)

    # Greeting (tight spacing)
    add_compact_paragraph(f"Dear {owner_name},", space_after=24)

    # Hook (tight spacing + bold)
    add_compact_paragraph(
        "Homeowners Insurance Rates Are Finally on the Decline – Don’t Miss Out!",
        bold=True,
        space_after=10
    )

    # Determine county name
    city_state = zip_to_city_state(zip_code)
    county_name = "Indian River"
    if "County" in city_state:
        county_name = city_state.split("County")[0].strip()

    # Letter body (normal spacing)
    body_text = (
        "For the first time in years, homeowners rates are coming down — and the savings could be significant.\n\n"
        "Recent legislative changes have boosted competition in Florida’s property insurance market, "
        f"and many {county_name} County homeowners are already benefiting.\n\n"
        "Jones Insurance Advisors is a two-generation, family-owned independent agency located right here in Vero Beach. "
        "Our team of dedicated agents possess extensive knowledge of the intricacies of the local insurance market, "
        "and are excited to assist you in finding the most comprehensive and competitively priced insurance solutions.\n\n"
        "Call us today for a free, no-obligation quote, or visit our website below and complete a quote request, "
        "and one of our dedicated agents will reach out to you!\n\n"
        "We look forward to earning your business and providing you the personal, dedicated service you have come to "
        "expect by doing business locally.\n\n"
        "Warm Regards,"
    )
    doc.add_paragraph(body_text)

    # Signature image
    if os.path.exists("signature_brian.png"):
        doc.add_picture("signature_brian.png", width=Inches(1.5), height=Inches(0.5))

    # Signature block
    doc.add_paragraph(
        f"{YOUR_NAME}\n{YOUR_TITLE}\n{YOUR_EMAIL}\n{YOUR_PHONE}\n{YOUR_WEB}"
    )

    doc.add_page_break()



# === ADD ENVELOPE TO DOC ===
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_envelope_to_doc(doc, owner_name, address, zip_code):
    # Add a new section for each envelope with correct size and margins
    section = doc.add_section()
    section.page_width = Inches(9.5)
    section.page_height = Inches(4.125)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)

    # Add return address paragraph - left aligned, small font
    sender = doc.add_paragraph(YOUR_RETURN_ADDRESS)
    sender.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in sender.runs:
        run.font.size = Pt(10)

    # Add a paragraph with specific spacing to push recipient block down
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(40)  # Adjust to move recipient block down

    # Recipient address block centered
    recipient = doc.add_paragraph()
    recipient.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Owner name bold, font size 14pt
    name_run = recipient.add_run(f"{owner_name}\n")
    name_run.bold = True
    name_run.font.size = Pt(14)

    # Address and city/state/zip normal font size 14pt
    addr_run = recipient.add_run(f"{address}\n{zip_to_city_state(zip_code)}")
    addr_run.font.size = Pt(14)

    # Add a page break to ensure next envelope starts fresh
    doc.add_page_break()




# === CREATE LABELS DOC ===
def create_labels(label_data):
    doc = Document()

    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.19)
    section.right_margin = Inches(0.19)

    labels_per_row = 3
    label_width = Inches(2.63)
    label_height = Inches(1.0)

    # Calculate how many rows you need
    num_rows = (len(label_data) + labels_per_row - 1) // labels_per_row

    table = doc.add_table(rows=num_rows, cols=labels_per_row)
    table.autofit = False
    table.allow_autofit = False

    for col in table.columns:
        col.width = label_width

    for row in table.rows:
        row.height = label_height
        row.height_rule = 2  # EXACTLY

    idx = 0
    for row in table.rows:
        for cell in row.cells:
            if idx < len(label_data):
                lines = label_data[idx].split("\n")
                para = cell.paragraphs[0]
                para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                if lines:
                    name_run = para.add_run(lines[0] + "\n")
                    name_run.bold = True
                    name_run.font.size = Pt(10.5)

                for line in lines[1:]:
                    addr_run = para.add_run(line + "\n")
                    addr_run.font.size = Pt(10.5)

                idx += 1
            else:
                cell.text = ""

    doc.save(LABELS_FILE)
    print(f"✅ Mailing labels saved to: {LABELS_FILE}")


# === MAIN ===
import time

def lookup_yellow_pages(name, city_state):
    import requests

    url = "https://yellow-pages-api.p.rapidapi.com/search"
    headers = {
        "X-RapidAPI-Key": "",
        "X-RapidAPI-Host": "yellow-pages3.p.rapidapi.com"
    }
    params = {
        "query": name,
        "location": city_state
    }

    try:
        response = requests.get(url, headers=headers, params=params)
        if response.status_code == 429:
            print(f"⚠️ Rate limit hit for {name}, retrying...")
            time.sleep(0)
            response = requests.get(url, headers=headers, params=params)  # retry once
        if response.status_code != 200:
            print(f"⚠️ Yellow Pages API error for {name}: {response.status_code}")
            return None 

        data = response.json()
        # Extract phone/email if available
        results = data.get("results", [])
        if results:
            first_result = results[0]
            phone = first_result.get("phone", "")
            emails = first_result.get("emails", [])
            email = emails[0] if emails else ""
            return {"phone": phone, "email": email}
        return None
    except Exception as e:
        print(f"⚠️ Yellow Pages API exception for {name}: {e}")
        return None

def main():
    load_zip_lookup()

    if not os.path.exists(EXCEL_FILE):
        print(f"❌ File not found: {EXCEL_FILE}")
        return

    try:
        df = pd.read_excel(EXCEL_FILE)
    except Exception as e:
        print(f"⚠️ Failed to read Excel file: {e}")
        return

    labels = []
    crm_rows = []

    for _, row in df.iterrows():
        try:
            raw_owner = str(row['Owner Name'])
            address = str(row['Address']).title().strip()
            zip_code = str(row['Site Zip Code']).strip()
            mailing_address = str(row['Mailing Address']).strip()
            sale_date_raw = str(row['Sale Date']).strip()
            sale_price_str = str(row['Sale Price']).replace('$', '').replace(',', '').strip()
            sale_price = float(sale_price_str)

            if not is_owner_occupied(address, mailing_address):
                print(f"⏭️ Skipping non-owner-occupied: {raw_owner}")
                continue

            owner_name = clean_owner_name(raw_owner)
            sale_date = datetime.strptime(sale_date_raw, '%m/%d/%Y').strftime('%B %d, %Y')

            # For Yellow Pages lookup, clean to just first and last name of first person only
            first_owner_part = raw_owner.split('||')[0].strip()
            # Remove suffixes and extra parentheses
            first_owner_part = first_owner_part.split('(')[0].strip()
            name_words = first_owner_part.split()
            if len(name_words) >= 2:
                lookup_name = f"{name_words[1]} {name_words[0]}"  # Usually First Last order
            elif len(name_words) == 1:
                lookup_name = name_words[0]
            else:
                lookup_name = owner_name  # fallback

            city_state = zip_to_city_state(zip_code)

            # Lookup phone and email from Yellow Pages API
            yp_info = lookup_yellow_pages(lookup_name, city_state)
            phone = yp_info['phone'] if yp_info and yp_info.get('phone') else ''
            email = yp_info['email'] if yp_info and yp_info.get('email') else ''

            add_letter_to_doc(letters_doc, owner_name, address, zip_code, sale_date, sale_price)
            add_envelope_to_doc(envelopes_doc, owner_name, address, zip_code)

            label_text = f"{owner_name}\n{address}\n{zip_to_city_state(zip_code)}"
            labels.append(label_text)

            crm_rows.append({
                'Name': lookup_name,   # Only first & last name for CRM
                'Address': address,
                'Zip': zip_code,
                'Sale Date': sale_date,
                'Sale Price': sale_price,
                'Email': email,
                'Phone': phone,
                'Source': 'Homeowner Anniversary Mailer-Sept-Oct'
            })

            print(f"✅ Processed: {owner_name} (YP Phone: {phone}, Email: {email})")

            time.sleep(1)  # Delay 1 second between API calls to avoid rate limit

        except Exception as e:
            print(f"⚠️ Skipped row due to error: {e}")

    letters_doc.save(LETTERS_FILE)
    envelopes_doc.save(ENVELOPES_FILE)
    print(f"📄 All letters saved to: {LETTERS_FILE}")
    print(f"✉️ All envelopes saved to: {ENVELOPES_FILE}")

    if labels:
        create_labels(labels)

    if crm_rows:
        keys = crm_rows[0].keys()
        with open(CRM_EXPORT_FILE, 'w', newline='', encoding='utf-8') as f:
            dict_writer = csv.DictWriter(f, keys)
            dict_writer.writeheader()
            dict_writer.writerows(crm_rows)
        print(f"📥 CRM-ready CSV saved to: {CRM_EXPORT_FILE}")

if __name__ == "__main__":
    main()



### Prints the credits in the terminal ###

def print_logo():
    logo = r"""
                      __/___             
             _____/______|           
     _______/_____\_______\_____     
     \              < < <       |    
      ~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~  
       AutoMailerPro v5.1
       Author: Kyle Padilla
       Created: 08/09/2025
       
       Version: 05.01
    """
    print(logo)

print_logo()
